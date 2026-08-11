"""Real format conversion — replaces the old `ConvertDialog` (mock `QTimer`,
never actually read/wrote any file, see history in `app/ui/modals.py`). Each
function takes an explicit source path + destination path, returns the
destination path on success, raises `ConvertError` (Vietnamese message) on
failure.

`ConvertError` inherits from `ApiError` so `app.thread_worker.CallWorker`
emits the right message automatically (without wrapping it in an "Unknown
error:" prefix) — see how `CallWorker.run()` distinguishes `ApiError` from
other errors.

Most functions run locally, no network needed — only `pdf_to_excel_file()`
needs an AI call (Claude reads the PDF directly via `/extract_table`), so it
needs network + quota.
"""
import base64
from pathlib import Path

from PIL import Image

from ...account import api_client
from ..excel.com import ExcelComError, export_workbook_pdf
from ..image import ops as image_ops
from ..pptx.render import ppt_to_pdf
from ..word.render import word_to_pdf
from .creators import create_excel_from_text
from .extractors import pdf_page_texts

from ...logging_config import get_logger, log_call
from ...strings import Converters as S

logger = get_logger(__name__)

_IMG_FORMAT_BY_EXT = {
    ".png": "PNG", ".jpg": "JPEG", ".jpeg": "JPEG",
    ".webp": "WEBP", ".bmp": "BMP",
}


class ConvertError(api_client.ApiError):
    """Conversion failed — message shown to the user in Vietnamese."""

    @log_call
    def __init__(self, message: str):
        super().__init__(message, "convert_failed")


# ── Word/Excel/PPT → PDF (COM) ────────────────────────────────────────────────

@log_call
def word_to_pdf_file(path: str, out_path: str) -> str:
    try:
        return word_to_pdf(path, out_path)
    except Exception as exc:  # noqa: BLE001 — all COM errors collapse to one message
        raise ConvertError(S.WORD_TO_PDF_FAILED.format(error=exc))


@log_call
def ppt_to_pdf_file(path: str, out_path: str) -> str:
    try:
        return ppt_to_pdf(path, out_path)
    except Exception as exc:  # noqa: BLE001
        raise ConvertError(S.PPT_TO_PDF_FAILED.format(error=exc))


@log_call
def excel_to_pdf_file(path: str, out_path: str) -> str:
    try:
        return export_workbook_pdf(path, out_path)
    except ExcelComError as exc:
        raise ConvertError(str(exc))


# ── PDF → Word/Excel/Image ────────────────────────────────────────────────────

@log_call
def pdf_to_word_file(path: str, out_path: str) -> str:
    """Recreates basic content: reads the text in the correct order, does NOT
    keep the original layout/style. Only works with PDFs that have a text
    layer — a pure scanned-image PDF should go through the image OCR flow
    (Image processing) instead of this function."""
    pages = pdf_page_texts(path)
    if not any(page_text.strip() for page_text in pages):
        raise ConvertError(S.PDF_NO_TEXT_LAYER)

    from docx import Document as DocxDocument
    doc = DocxDocument()
    for page_index, text in enumerate(pages):
        if page_index > 0:
            doc.add_page_break()
        lines = [line for line in text.split("\n") if line.strip()] or ["(trang trống)"]
        for line in lines:
            doc.add_paragraph(line)
    try:
        doc.save(out_path)
    except PermissionError:
        raise ConvertError(S.WRITE_FAILED_FILE_OPEN.format(name=Path(out_path).name))
    return out_path


@log_call
def pdf_to_excel_file(path: str, out_path: str) -> str:
    """Have Claude read the PDF directly (even a scanned copy) then extract
    the table — needs network + quota. Reuses the same "--- Sheet: X ---" +
    CSV convention that `create_excel_from_text()` already parses for the
    "create a new document" flow."""
    data = Path(path).read_bytes()
    attachment = {
        "kind": "pdf", "media_type": "application/pdf",
        "data_b64": base64.b64encode(data).decode("ascii"),
    }
    result = api_client.extract_table(attachment)
    text = (result.get("text") or "").strip()
    if not text:
        raise ConvertError(S.NO_TABLE_DATA_FROM_FILE)
    create_excel_from_text(text, out_path)
    return out_path


@log_call
def pdf_to_images(path: str, out_dir: str) -> list[str]:
    import fitz
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    doc = fitz.open(path)
    out_paths: list[str] = []
    try:
        for page_index, page in enumerate(doc):
            pix = page.get_pixmap(dpi=150)
            page_path = str(Path(out_dir) / f"trang_{page_index + 1:03d}.png")
            pix.save(page_path)
            out_paths.append(page_path)
    finally:
        doc.close()
    if not out_paths:
        raise ConvertError(S.PDF_NO_PAGES)
    return out_paths


# ── Image → PDF / format conversion ─────────────────────────────────────────

@log_call
def image_to_pdf_file(paths: list[str], out_path: str) -> str:
    if not paths:
        raise ConvertError(S.NO_IMAGES_TO_CONVERT)
    images = []
    for image_path in paths:
        img = Image.open(image_path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)
    try:
        images[0].save(out_path, "PDF", save_all=True, append_images=images[1:])
    except Exception as exc:  # noqa: BLE001
        raise ConvertError(S.IMAGE_TO_PDF_FAILED.format(error=exc))
    return out_path


@log_call
def image_to_word_file(paths: list[str], out_path: str) -> str:
    """OCR each image via Claude Vision (`/extract_text`) then build 1 editable
    Word file — each image its own page (page break between images), keeping
    the paragraph breaks the AI read. Needs network + quota, costs 1 call/image."""
    if not paths:
        raise ConvertError(S.NO_IMAGES_TO_CONVERT)
    from docx import Document as DocxDocument
    doc = DocxDocument()
    got_any = False
    for page_index, path in enumerate(paths):
        if page_index > 0:
            doc.add_page_break()
        data_b64, media_type = image_ops.b64_for_vision(path)
        attachment = {"kind": "image", "media_type": media_type, "data_b64": data_b64}
        result = api_client.extract_text(attachment)
        text = (result.get("text") or "").strip()
        if text and "Khong doc duoc" not in text:
            got_any = True
        for para in (text or "(không đọc được nội dung)").split("\n\n"):
            doc.add_paragraph(para.strip())
    if not got_any:
        raise ConvertError(S.NO_TEXT_READ_FROM_IMAGES)
    try:
        doc.save(out_path)
    except PermissionError:
        raise ConvertError(S.WRITE_FAILED_FILE_OPEN.format(name=Path(out_path).name))
    return out_path


@log_call
def image_to_searchable_pdf_file(paths: list[str], out_path: str) -> str:
    """Export a PDF from photos with a hidden OCR text layer (`render_mode=3`
    — not drawn but still selectable/searchable/copyable) — unlike
    `image_to_pdf_file()` which only embeds the raw image. The text layer
    covers the whole page, does NOT match each character's exact position
    since Claude Vision returns plain text with no coordinates — still
    enough for full-text search/copy. Needs network + quota, costs 1
    call/image."""
    if not paths:
        raise ConvertError(S.NO_IMAGES_TO_CONVERT)
    import fitz
    doc = fitz.open()
    try:
        for path in paths:
            img = Image.open(path)
            width, height = img.size
            page = doc.new_page(width=width, height=height)
            page.insert_image(page.rect, filename=path)

            data_b64, media_type = image_ops.b64_for_vision(path)
            attachment = {"kind": "image", "media_type": media_type, "data_b64": data_b64}
            result = api_client.extract_text(attachment)
            text = (result.get("text") or "").strip()
            if text and "Khong doc duoc" not in text:
                page.insert_textbox(page.rect, text, fontsize=8, render_mode=3)
        doc.save(out_path)
    except Exception as exc:  # noqa: BLE001
        raise ConvertError(S.SEARCHABLE_PDF_FAILED.format(error=exc))
    finally:
        doc.close()
    return out_path


@log_call
def image_to_excel_file(path: str, out_path: str) -> str:
    """Have Claude read the image (photographed invoice/contract) via Vision
    then extract the table — needs network + quota. Same /extract_table and
    same "--- Sheet: X ---" + CSV convention as pdf_to_excel_file()."""
    data_b64, media_type = image_ops.b64_for_vision(path)
    attachment = {"kind": "image", "media_type": media_type, "data_b64": data_b64}
    result = api_client.extract_table(attachment)
    text = (result.get("text") or "").strip()
    if not text:
        raise ConvertError(S.NO_TABLE_DATA_FROM_IMAGE)
    create_excel_from_text(text, out_path)
    return out_path


@log_call
def image_convert_format(path: str, out_path: str) -> str:
    target_fmt = _IMG_FORMAT_BY_EXT.get(Path(out_path).suffix.lower())
    if target_fmt is None:
        raise ConvertError(S.TARGET_FORMAT_UNSUPPORTED.format(ext=Path(out_path).suffix))
    try:
        img = Image.open(path)
        if target_fmt == "JPEG" and img.mode != "RGB":
            img = img.convert("RGB")
        img.save(out_path, target_fmt)
    except Exception as exc:  # noqa: BLE001
        raise ConvertError(S.IMAGE_FORMAT_CONVERT_FAILED.format(error=exc))
    return out_path


# ── Common entry point for ConvertDialog ─────────────────────────────────────

@log_call
def convert_file(path: str, source_type: str, target_ext: str, out_path: str,
                  extra_paths: list[str] | None = None, searchable: bool = False) -> str:
    """`source_type`: "word"|"excel"|"ppt"|"pdf"|"image". Returns the result
    path — for PDF→Image, `out_path` is a FOLDER holding the PNG pages.

    `extra_paths` (only applies when source_type="image", target .pdf/.docx):
    additional images besides `path` to merge into a single multi-page file,
    in order. `searchable` (only applies when target is .pdf): True → PDF
    with a hidden OCR text layer (`image_to_searchable_pdf_file`), False →
    embeds the raw image as before.
    """
    if source_type in ("word", "excel", "ppt") and target_ext == ".pdf":
        fn = {"word": word_to_pdf_file, "excel": excel_to_pdf_file, "ppt": ppt_to_pdf_file}[source_type]
        return fn(path, out_path)
    if source_type == "pdf" and target_ext == ".docx":
        return pdf_to_word_file(path, out_path)
    if source_type == "pdf" and target_ext == ".xlsx":
        return pdf_to_excel_file(path, out_path)
    if source_type == "pdf" and target_ext == ".png":
        pdf_to_images(path, out_path)
        return out_path
    if source_type == "image" and target_ext == ".pdf":
        all_paths = [path] + (extra_paths or [])
        if searchable:
            return image_to_searchable_pdf_file(all_paths, out_path)
        return image_to_pdf_file(all_paths, out_path)
    if source_type == "image" and target_ext == ".docx":
        return image_to_word_file([path] + (extra_paths or []), out_path)
    if source_type == "image" and target_ext == ".xlsx":
        return image_to_excel_file(path, out_path)
    if source_type == "image" and target_ext in _IMG_FORMAT_BY_EXT:
        return image_convert_format(path, out_path)
    raise ConvertError(S.CONVERSION_UNSUPPORTED.format(source_type=source_type, target_ext=target_ext))
