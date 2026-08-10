import csv
import io
import re

from ...logging_config import get_logger, log_call

logger = get_logger(__name__)

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
except ImportError:
    logger.warning("python-docx not installed — save_word() will fail if called.")
    DocxDocument = None  # type: ignore
    Pt = None            # type: ignore

try:
    import openpyxl
except ImportError:
    logger.warning("openpyxl not installed — save_excel() will fail if called.")
    openpyxl = None  # type: ignore


@log_call
def save_word(original_path: str, edited_text: str, save_path: str):
    doc = DocxDocument()
    try:
        orig = DocxDocument(original_path)
        doc.styles["Normal"].font.name = orig.styles["Normal"].font.name or "Calibri"
        doc.styles["Normal"].font.size = orig.styles["Normal"].font.size or Pt(11)
    except Exception:
        logger.debug("Could not copy base style from original doc %s — using default.",
                     original_path, exc_info=True)
    for line in edited_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(save_path)


@log_call
def save_excel(edited_text: str, save_path: str):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    ws = None
    for line in edited_text.splitlines():
        match = re.match(r"^---\s*Sheet:\s*(.+?)\s*---\s*$", line)
        if match:
            ws = wb.create_sheet(title=match.group(1)[:31])
            continue
        if ws is None:
            ws = wb.create_sheet(title="Sheet1")
        if line.strip():
            for row in csv.reader(io.StringIO(line)):
                ws.append(row)
    wb.save(save_path)
