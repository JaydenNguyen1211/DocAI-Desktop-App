"""The aggregate function for Folder mode — takes a list of files (Word/
Excel/PowerPoint/PDF/Image) plus the user's request, returns analysis
content (always present) and a list of output files (may be empty), per the
design flow at `D:\\Tools\\DocAI\\Plan\\V4\\DocAI-dafiile\\screens`.

6 tasks are recognized locally (no AI call spent on classification):
    compare · reconcile · merge · extract & merge mixed sources · search ·
    review for missing documents
Requests that don't match any of the above are passed verbatim to
`api_client.chat()` — this is the only "escalate" point that currently
exists (the server picks the right model itself; the client has no model
selection parameter).

Output files are written to the temp staging folder (`_staging_dir()`), NOT
the final save location yet — per the "only create the real file when Save
file is clicked" behavior in the design. Opening the save-location dialog is
a separate UI step (see `app/ui/save_output.py`).
"""
import copy
import os
import re
import tempfile
import time
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

from ...account import api_client
from ..image import ops as image_ops
from . import extractors
from .creators import create_excel_from_text
from .models import AttachedFile

from ...logging_config import get_logger, log_call
from ...strings import DocSet as S

logger = get_logger(__name__)

TASK_COMPARE = "compare"
TASK_RECONCILE = "reconcile"
TASK_MERGE = "merge"
TASK_EXTRACT_MERGE = "extract_merge"
TASK_SEARCH = "search"
TASK_CHECKLIST = "checklist"
TASK_GENERAL = "general"

_FAMILY_BY_EXT = {
    ".docx": "word", ".doc": "word",
    ".xlsx": "excel", ".xls": "excel",
    ".pptx": "ppt", ".ppt": "ppt",
    ".pdf": "pdf",
    ".png": "image", ".jpg": "image", ".jpeg": "image",
    ".bmp": "image", ".webp": "image", ".gif": "image",
}

_MAX_CHARS_PER_FILE = 6000


class DocSetError(api_client.ApiError):
    """Couldn't process the file set — message shown to the user in Vietnamese."""

    @log_call
    def __init__(self, message: str):
        super().__init__(message, "doc_set_failed")


@dataclass
class DocSetResult:
    task: str
    analysis: str
    output_files: list[str] = field(default_factory=list)
    quota: dict = field(default_factory=dict)


# ── Task classification (local, no AI call) ─────────────────────────────────

_RE_MERGE = re.compile(r"\bgộp\b|\bghép\b|hợp nhất thành|\bmerge\b", re.IGNORECASE)
_RE_CHECKLIST = re.compile(
    r"còn thiếu|thiếu gì|đầy đủ chưa|rà soát|checklist|kiểm tra hồ sơ|thiếu giấy tờ",
    re.IGNORECASE)
_RE_SEARCH = re.compile(r"\btìm\b|\btra cứu\b|\bsearch\b", re.IGNORECASE)
_RE_RECONCILE = re.compile(
    r"đối chiếu|có khớp|khớp không|chênh lệch|sai lệch|trùng khớp", re.IGNORECASE)
_RE_COMPARE = re.compile(r"so sánh|khác biệt|khác nhau|\bcompare\b", re.IGNORECASE)
_RE_EXTRACT = re.compile(r"trích xuất|tổng hợp|\bextract\b", re.IGNORECASE)


@log_call
def _classify(user_request: str, paths: list[str]) -> str:
    text = user_request or ""
    families = {_FAMILY_BY_EXT.get(os.path.splitext(p)[1].lower()) for p in paths}
    same_family = len(families) == 1 and None not in families

    if _RE_MERGE.search(text):
        return TASK_MERGE if same_family else TASK_EXTRACT_MERGE
    if _RE_CHECKLIST.search(text):
        return TASK_CHECKLIST
    if _RE_SEARCH.search(text):
        return TASK_SEARCH
    if _RE_RECONCILE.search(text):
        return TASK_RECONCILE
    if _RE_COMPARE.search(text):
        return TASK_COMPARE
    if _RE_EXTRACT.search(text) and not same_family:
        return TASK_EXTRACT_MERGE
    return TASK_GENERAL


# ── Determine input files from the chat context (local, no AI call) ────────
#
# The Chat Panel is the center — always opens independently, no sidebar
# action needed first. The sidebar (folder tree) is only a REFERENCE SOURCE
# to recognize file names mentioned in the chat message, there's no more
# "tick to select" concept.
#
# 3 candidate sources (all only "COULD BE input", never auto-used):
#   - `attached`: files attached via chat in this conversation
#   - `recent_outputs`: output files from a previous multi-file run in the
#     SAME conversation
#   - `folder_files`: every file in the currently open folder (only used to
#     recognize names mentioned in the chat, or when the user clearly means
#     "the whole folder")
# Which files are ACTUALLY the input is decided by what the user explicitly
# said (named files, or "the whole folder" intent) — sources are never
# merged together automatically. If it's not clear enough, return
# `needs_clarification=True` with suggestions to ask back instead of guessing.

_WHOLE_FOLDER_RE = re.compile(
    r"toàn bộ thư mục|tất cả file|cả thư mục|mọi file|hết các file|toàn bộ file",
    re.IGNORECASE)
_RECENT_OUTPUT_RE = re.compile(
    r"vừa (gộp|tạo|xuất|xong)|kết quả (vừa rồi|trên|này)|file (vừa|mới) (tạo|gộp|xuất)",
    re.IGNORECASE)
_COUNT_WORDS = {"hai": 2, "2": 2, "ba": 3, "3": 3, "bốn": 4, "4": 4, "năm": 5, "5": 5}
_COUNT_HINT_RE = re.compile(r"\b(hai|ba|bốn|năm|\d+)\s+file\b", re.IGNORECASE)


@dataclass
class ResolveResult:
    paths: list[str]
    needs_clarification: bool = False
    message: str = ""
    suggestions: list[str] = field(default_factory=list)  # suggested file names (chips)


@log_call
def _name_mentioned(text_words: set[str], path: str, threshold: float = 0.6) -> bool:
    """Match by WORD-overlap ratio, not a hard substring — a file name like
    "HopDong_LoHangT6.docx" (splits into "hop dong lo hang t6") needs to
    match the natural sentence "hợp đồng ... lô hàng T6", which doesn't
    literally contain the file name."""
    stem_words = {w for w in _norm(os.path.splitext(os.path.basename(path))[0]).split() if len(w) >= 2}
    if not stem_words:
        return False
    return len(stem_words & text_words) / len(stem_words) >= threshold


@log_call
def resolve_input_files(user_request: str, attached: list[str], recent_outputs: list[str],
                        folder_files: list[str]) -> ResolveResult:
    text = user_request or ""
    text_words = {w for w in _norm(text).split() if len(w) >= 2}

    pool: list[str] = []
    for group in (attached, recent_outputs, folder_files):
        for path in group:
            if path not in pool:
                pool.append(path)

    named = [path for path in pool if _name_mentioned(text_words, path)]
    if _RECENT_OUTPUT_RE.search(text) and recent_outputs and recent_outputs[-1] not in named:
        named.append(recent_outputs[-1])

    looks_like_wholeset = bool(_RE_SEARCH.search(text) or _RE_CHECKLIST.search(text))

    if named:
        resolved = named
    elif _WHOLE_FOLDER_RE.search(text) and folder_files:
        resolved = list(folder_files)
    elif attached:
        resolved = list(attached)
    elif recent_outputs and _RECENT_OUTPUT_RE.search(text):
        resolved = [recent_outputs[-1]]
    elif looks_like_wholeset and folder_files:
        resolved = list(folder_files)
    else:
        resolved = []

    if not resolved:
        return ResolveResult(
            paths=[], needs_clarification=True,
            message=S.CLARIFY_WHICH_FILES,
            suggestions=[os.path.basename(p) for p in pool[:8]])

    hint_match = _COUNT_HINT_RE.search(text)
    if hint_match:
        hint_count = _COUNT_WORDS.get(hint_match.group(1).lower(), 0)
        if hint_count and len(resolved) < hint_count:
            remaining = [path for path in pool if path not in resolved]
            return ResolveResult(
                paths=resolved, needs_clarification=True,
                message=S.CLARIFY_MORE_FILES.format(hint_count=hint_count, resolved_count=len(resolved)),
                suggestions=[os.path.basename(p) for p in remaining[:8]])

    return ResolveResult(paths=resolved)


# ── Extract content from every file type ────────────────────────────────────

@log_call
def _extract_one(path: str) -> AttachedFile:
    ext = os.path.splitext(path)[1].lower()
    family = _FAMILY_BY_EXT.get(ext)
    if family == "word":
        return extractors.extract_word(path)
    if family == "excel":
        return extractors.extract_excel(path)
    if family == "ppt":
        return extractors.extract_pptx(path)
    if family == "pdf":
        pages = extractors.pdf_page_texts(path)
        if any(page_text.strip() for page_text in pages):
            return extractors.extract_pdf(path)
        return _ocr_pdf(path)
    if family == "image":
        return _ocr_image(path)
    raise DocSetError(S.FORMAT_UNSUPPORTED.format(ext_or_path=ext or path))


@log_call
def _extract_all(paths: list[str]) -> tuple[list[AttachedFile], list[str]]:
    """Read the content of each file — 1 failed file (odd format, corrupted…)
    does NOT break the whole batch, especially important when the input is
    "the whole folder" (e.g. an e-invoice .xml file mixed into the folder
    with no dedicated extractor yet). Returns (readable files, list of notes
    about skipped files)."""
    files: list[AttachedFile] = []
    skipped: list[str] = []
    for path in paths:
        try:
            files.append(_extract_one(path))
        except DocSetError as exc:
            logger.warning("Skipping file in batch extract: path=%s reason=%s",
                           path, exc.message)
            skipped.append(f"{os.path.basename(path)} ({exc.message})")
    return files, skipped


@log_call
def _ocr_pdf(path: str) -> AttachedFile:
    """A scanned-image PDF (no text layer) — have the AI read it via Vision.
    Costs 1 AI action."""
    import base64
    data = Path(path).read_bytes()
    attachment = {
        "kind": "pdf", "media_type": "application/pdf",
        "data_b64": base64.b64encode(data).decode("ascii"),
    }
    result = api_client.extract_text(attachment)
    text = (result.get("text") or "").strip()
    return AttachedFile(
        path=path, name=os.path.basename(path), file_type="pdf",
        claude_content=text or "(không đọc được nội dung — PDF quét ảnh chất lượng thấp)")


@log_call
def _ocr_image(path: str) -> AttachedFile:
    """Costs 1 AI action / image — there's no way to read text in an image locally."""
    data_b64, media_type = image_ops.b64_for_vision(path)
    attachment = {"kind": "image", "media_type": media_type, "data_b64": data_b64}
    result = api_client.extract_text(attachment)
    text = (result.get("text") or "").strip()
    return AttachedFile(
        path=path, name=os.path.basename(path), file_type="image",
        claude_content=text or "(không đọc được nội dung ảnh)")


@log_call
def _combine_for_prompt(files: list[AttachedFile]) -> str:
    parts = []
    for index, attached in enumerate(files, start=1):
        content = attached.claude_content or ""
        if len(content) > _MAX_CHARS_PER_FILE:
            content = content[:_MAX_CHARS_PER_FILE] + "\n… (đã cắt bớt, file quá dài)"
        parts.append(f"=== File {index}: {attached.name} ===\n{content}")
    return "\n\n".join(parts)


@log_call
def _staging_dir() -> str:
    out_dir = os.path.join(tempfile.gettempdir(), "DocAI", "staging")
    os.makedirs(out_dir, exist_ok=True)
    return out_dir


@log_call
def _stamp() -> str:
    return time.strftime("%Y%m%d_%H%M%S")


@log_call
def _quota_of(result: dict) -> dict:
    """`plan`/`quota_remaining` from a server response — lets the UI update
    the quota label just like the regular chat/edit flow (the server returns
    the current quota state, not this call's own consumption)."""
    quota = {}
    if result.get("quota_remaining") is not None:
        quota["quota_remaining"] = result["quota_remaining"]
    if result.get("plan") is not None:
        quota["plan"] = result["plan"]
    return quota


# ── Cross-file search (local, no AI call) ────────────────────────────────────

_QUOTE_RE = re.compile(r'["“”\'‘’](.+?)["“”\'‘’]')
_SEARCH_LEAD_RE = re.compile(
    r'^(tìm kiếm|tìm giúp|tìm|tra cứu|search)\s+(từ khóa\s+)?', re.IGNORECASE)
_SEARCH_TRAIL_RE = re.compile(
    r'\btrong\s+(toàn bộ\s+|cả\s+)?(thư mục|các file|tất cả file|hồ sơ).*$',
    re.IGNORECASE)


@log_call
def _extract_search_keyword(user_request: str) -> str:
    text = (user_request or "").strip()
    quoted = _QUOTE_RE.search(text)
    if quoted:
        return quoted.group(1).strip()
    cleaned = _SEARCH_LEAD_RE.sub("", text)
    cleaned = _SEARCH_TRAIL_RE.sub("", cleaned)
    return cleaned.strip()


@log_call
def _find_snippets(content: str, keyword: str, context: int = 40) -> list[str]:
    if not keyword or not content:
        return []
    lower_content = content.lower()
    lower_keyword = keyword.lower()
    out = []
    start = 0
    while True:
        idx = lower_content.find(lower_keyword, start)
        if idx == -1:
            break
        span_start = max(0, idx - context)
        span_end = min(len(content), idx + len(keyword) + context)
        out.append(content[span_start:span_end].replace("\n", " ").strip())
        start = idx + len(keyword)
    return out


@log_call
def _handle_search(files: list[AttachedFile], user_request: str) -> DocSetResult:
    keyword = _extract_search_keyword(user_request)
    if not keyword:
        return DocSetResult(
            task=TASK_SEARCH,
            analysis='Chưa xác định được từ khóa cần tìm — hãy nói rõ hơn, '
                     'VD: Tìm "điều khoản phạt vi phạm" trong các file này.')

    per_file: list[tuple[AttachedFile, list[str]]] = []
    total_matches = 0
    for attached in files:
        matches = _find_snippets(attached.claude_content, keyword)
        if matches:
            per_file.append((attached, matches))
            total_matches += len(matches)

    if not per_file:
        analysis = f'Không tìm thấy "{keyword}" trong {len(files)} file đã chọn.'
        return DocSetResult(task=TASK_SEARCH, analysis=analysis)

    lines = [f'Tìm thấy {total_matches} kết quả cho "{keyword}" trong {len(per_file)} file:\n']
    for attached, matches in per_file:
        lines.append(f"**{attached.name}** — {len(matches)} lần xuất hiện:")
        for snippet in matches[:5]:
            lines.append(f"  … {snippet} …")
        if len(matches) > 5:
            lines.append(f"  … và {len(matches) - 5} lần khác")
    return DocSetResult(task=TASK_SEARCH, analysis="\n".join(lines))


# ── Review for missing documents (mostly local, 1 AI call if the checklist
#    needs to be inferred) ────────────────────────────────────────────────

_STOPWORDS = {"giay", "to", "ho", "so", "va", "cac", "mot", "cho", "cua", "theo", "cac"}


_CAMEL_RE = re.compile(r"(?<=[a-z0-9])(?=[A-Z])")


@log_call
def _norm(text: str) -> str:
    """Normalize for matching: split camelCase (e.g. "HopDong" → "Hop Dong" —
    file names are often concatenated this way, no space/underscore between
    words), strip Vietnamese diacritics, lowercase, collapse every
    non-alphanumeric character into a single space."""
    text = _CAMEL_RE.sub(" ", text)
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^a-zA-Z0-9]+", " ", text.lower())
    return text


@log_call
def _infer_checklist(user_request: str, file_names: list[str]) -> list[str]:
    prompt = (
        "Người dùng muốn rà soát một bộ hồ sơ xem còn thiếu giấy tờ gì.\n"
        f"Yêu cầu của người dùng: {user_request}\n"
        f"Các file hiện có trong hồ sơ: {', '.join(file_names) or '(chưa có file nào)'}\n"
        "Hãy liệt kê danh mục các loại giấy tờ/tài liệu THƯỜNG CẦN CÓ cho loại hồ sơ này, "
        "mỗi mục một dòng, không đánh số, không giải thích thêm."
    )
    try:
        result = api_client.chat(prompt)
    except api_client.ApiError:
        logger.warning("Checklist inference AI call failed — falling back to empty checklist",
                       exc_info=True)
        return []
    text = (result.get("text") or "").strip()
    items = [line.strip("-•* \t") for line in text.splitlines() if line.strip()]
    return items[:15]


@log_call
def _match_checklist_item(item: str, files: list[AttachedFile]) -> AttachedFile | None:
    item_words = {w for w in _norm(item).split() if len(w) > 2 and w not in _STOPWORDS}
    if not item_words:
        return None
    best_file, best_score = None, 0.0
    for attached in files:
        haystack = _norm(os.path.splitext(attached.name)[0] + " " + attached.claude_content[:2000])
        haystack_words = set(haystack.split())
        score = len(item_words & haystack_words) / len(item_words)
        if score > best_score:
            best_file, best_score = attached, score
    return best_file if best_score >= 0.5 else None


@log_call
def _handle_checklist(files: list[AttachedFile], user_request: str,
                       checklist: list[str] | None) -> DocSetResult:
    items = checklist or _infer_checklist(user_request, [attached.name for attached in files])
    if not items:
        return DocSetResult(
            task=TASK_CHECKLIST,
            analysis="Chưa xác định được danh mục cần rà soát — hãy cung cấp checklist cụ thể "
                     'hoặc mô tả rõ loại hồ sơ (VD: "hồ sơ đấu thầu").')

    found_count = 0
    detail_lines = []
    for item in items:
        match = _match_checklist_item(item, files)
        if match:
            found_count += 1
            detail_lines.append(f"✓ {item} — {match.name}")
        else:
            detail_lines.append(f"✕ {item} — còn thiếu")

    header = f"Đủ {found_count}/{len(items)} mục.\n"
    analysis = header + "\n".join(detail_lines)
    analysis += ("\n\nLưu ý: đối chiếu theo tên file + trích đoạn nội dung — nên kiểm tra lại "
                 "thủ công với các mục ghi \"còn thiếu\" trước khi kết luận.")
    return DocSetResult(task=TASK_CHECKLIST, analysis=analysis)


# ── Compare / reconcile (AI analysis, text only — no file generated) ───────

@log_call
def _handle_compare(files: list[AttachedFile], user_request: str, task: str) -> DocSetResult:
    combined = _combine_for_prompt(files)
    if task == TASK_RECONCILE:
        verb = "đối chiếu số liệu/thông tin giữa các file, chỉ rõ nội dung nào khớp và nội dung nào lệch"
    else:
        verb = "so sánh nội dung giữa các file, liệt kê điểm khác biệt"
    prompt = (
        f"{user_request}\n\n"
        f"Hãy {verb}. Trả lời ngắn gọn dạng danh sách, mỗi ý một dòng bắt đầu bằng dấu '-', "
        "nêu rõ tên trường/nội dung và giá trị tương ứng ở từng file.\n\n"
        f"{combined}"
    )
    try:
        result = api_client.chat(prompt)
    except api_client.ApiError as exc:
        raise DocSetError(str(exc)) from exc
    text = (result.get("text") or "").strip() or "AI không đưa ra được kết quả so sánh."
    return DocSetResult(task=task, analysis=text, quota=_quota_of(result))


# ── Extract & merge mixed sources → 1 Excel table ───────────────────────────

_SHEET_BLOCK_RE = re.compile(r'(---\s*Sheet:.*)', re.DOTALL | re.IGNORECASE)


@log_call
def _split_sheet_block(text: str) -> tuple[str, str]:
    """Returns (analysis_before_the_table, csv_table_block) — the table block
    is empty if the AI doesn't return the "--- Sheet: X ---" format correctly."""
    match = _SHEET_BLOCK_RE.search(text)
    if not match:
        return text.strip(), ""
    return text[:match.start()].strip(), match.group(1).strip()


@log_call
def _handle_extract_merge(user_request: str, files: list[AttachedFile]) -> DocSetResult:
    combined = _combine_for_prompt(files)
    prompt = (
        f"{user_request}\n\n"
        "Dưới đây là nội dung trích xuất từ nhiều file nguồn khác định dạng. Hãy hợp nhất dữ "
        "liệu liên quan thành MỘT bảng duy nhất, khớp các trường theo tên/nhãn tương ứng giữa "
        "các nguồn. Nếu phát hiện chỗ vênh giữa các nguồn (VD cùng 1 đối tượng nhưng số liệu "
        "khác nhau), hãy nêu rõ TRƯỚC khi đưa ra bảng. Bảng trả về đúng định dạng:\n"
        "--- Sheet: <tên bảng> ---\n"
        "<dữ liệu CSV, dòng đầu là tiêu đề cột>\n\n"
        f"{combined}"
    )
    try:
        result = api_client.chat(prompt)
    except api_client.ApiError as exc:
        raise DocSetError(str(exc)) from exc
    text = (result.get("text") or "").strip()
    analysis, csv_block = _split_sheet_block(text)
    if not csv_block:
        return DocSetResult(
            task=TASK_EXTRACT_MERGE,
            analysis=text or "AI không tổng hợp được dữ liệu dạng bảng từ các nguồn này.")

    out_path = os.path.join(_staging_dir(), f"TongHop_{_stamp()}.xlsx")
    create_excel_from_text(csv_block, out_path)
    if not analysis:
        analysis = f"Đã hợp nhất dữ liệu từ {len(files)} nguồn thành 1 bảng."
    analysis += f"\n\nFile kết quả (bản xem trước, chưa lưu): {os.path.basename(out_path)}"
    return DocSetResult(task=TASK_EXTRACT_MERGE, analysis=analysis, output_files=[out_path],
                        quota=_quota_of(result))


# ── Merge same-format files (local, no AI call) ──────────────────────────────

@log_call
def _merge_word(paths: list[str], out_path: str) -> str:
    from docx import Document as DocxDocument
    merged = DocxDocument()
    merged.add_heading("Mục lục", level=1)
    for index, path in enumerate(paths, start=1):
        merged.add_paragraph(f"{index}. {os.path.splitext(os.path.basename(path))[0]}")
    merged.add_page_break()
    for file_index, path in enumerate(paths):
        if file_index > 0:
            merged.add_page_break()
        source = DocxDocument(path)
        for element in source.element.body:
            tag = element.tag.split("}")[-1]
            if tag == "sectPr":
                continue
            merged.element.body.append(copy.deepcopy(element))
    merged.save(out_path)
    return out_path


@log_call
def _merge_excel(paths: list[str], out_path: str) -> str:
    import openpyxl
    dst = openpyxl.Workbook()
    dst.remove(dst.active)
    used_names: set[str] = set()
    for path in paths:
        source = openpyxl.load_workbook(path, data_only=True)
        stem = os.path.splitext(os.path.basename(path))[0]
        for sheet_name in source.sheetnames:
            title = f"{stem}_{sheet_name}" if len(paths) > 1 else sheet_name
            title = title[:31]
            candidate, suffix = title, 1
            while candidate in used_names:
                suffix += 1
                candidate = f"{title[:28]}_{suffix}"
            used_names.add(candidate)
            dst_ws = dst.create_sheet(title=candidate)
            for row in source[sheet_name].iter_rows(values_only=True):
                dst_ws.append(list(row))
    dst.save(out_path)
    return out_path


@log_call
def _merge_pdf(paths: list[str], out_path: str) -> str:
    import fitz
    merged = fitz.open()
    try:
        for path in paths:
            with fitz.open(path) as source:
                merged.insert_pdf(source)
        merged.save(out_path)
    finally:
        merged.close()
    return out_path


_MERGE_BUILDERS = {"word": _merge_word, "excel": _merge_excel, "pdf": _merge_pdf}
_MERGE_EXT = {"word": ".docx", "excel": ".xlsx", "pdf": ".pdf"}


@log_call
def _handle_merge(paths: list[str], user_request: str) -> DocSetResult:
    families = {_FAMILY_BY_EXT.get(os.path.splitext(p)[1].lower()) for p in paths}
    if len(families) != 1 or None in families:
        # "merge" with mixed sources that share no common format -> the
        # original files can't be concatenated, redirect to merging the
        # data into a single table instead.
        files, skipped = _extract_all(paths)
        if not files:
            raise DocSetError(S.NO_READABLE_CONTENT)
        result = _handle_extract_merge(user_request, files)
        if skipped:
            result.analysis = (f"⚠ Bỏ qua {len(skipped)} file không đọc được: "
                               + ", ".join(skipped) + "\n\n" + result.analysis)
        return result

    family = next(iter(families))
    builder = _MERGE_BUILDERS.get(family)
    if builder is None:
        return DocSetResult(
            task=TASK_MERGE,
            analysis=f'Chưa hỗ trợ gộp trực tiếp file loại "{family}" ở đây. Với ảnh, dùng '
                     'tính năng "Gộp nhiều ảnh thành 1 file" ở màn hình Chuyển đổi.')

    out_path = os.path.join(_staging_dir(), f"Gop_{_stamp()}{_MERGE_EXT[family]}")
    try:
        builder(paths, out_path)
    except Exception as exc:  # noqa: BLE001 — all read/write errors collapse to one message
        raise DocSetError(S.MERGE_FAILED.format(error=exc)) from exc

    names = [os.path.basename(p) for p in paths]
    ordered = "\n".join(f"{i}. {name}" for i, name in enumerate(names, start=1))
    analysis = (f"Đã gộp {len(paths)} file theo đúng thứ tự đã chọn:\n{ordered}\n\n"
               f"File kết quả (bản xem trước, chưa lưu): {os.path.basename(out_path)}")
    return DocSetResult(task=TASK_MERGE, analysis=analysis, output_files=[out_path])


# ── Outside the 6 tasks above → hand off to chat() for general handling ────

@log_call
def _handle_general(files: list[AttachedFile], user_request: str,
                    history: list | None, business: dict | None) -> DocSetResult:
    combined = _combine_for_prompt(files)
    message = f"{user_request}\n\n{combined}" if files else user_request
    try:
        result = api_client.chat(message, business=business, history=history)
    except api_client.ApiError as exc:
        raise DocSetError(str(exc)) from exc
    text = (result.get("text") or "").strip() or "AI không có phản hồi."
    return DocSetResult(task=TASK_GENERAL, analysis=text, quota=_quota_of(result))


# ── Main entry point ──────────────────────────────────────────────────────────

@log_call
def process_document_set(paths: list[str], user_request: str,
                         checklist: list[str] | None = None,
                         history: list | None = None,
                         business: dict | None = None) -> DocSetResult:
    """The aggregate task for Folder mode.

    `paths`: list of input file paths, any type (Word/Excel/PowerPoint/PDF/
    Image). `user_request`: the command the user typed in chat. `checklist`:
    list of required item names, used only for the missing-documents review
    task — if left empty, the AI infers a reasonable checklist itself
    (costs 1 AI action).

    Returns `DocSetResult(task, analysis, output_files)` — `analysis` always
    has a value; `output_files` is only non-empty for the merge / extract &
    merge mixed sources tasks (the file is written to a temp folder, not the
    final save location yet).

    Cost note: image files and scanned-image PDFs (no text layer) always
    need 1 AI call per file to read their content (OCR), even for "local"
    tasks like search/review — there's no way to read text in an image
    without going through AI.
    """
    if not paths:
        raise DocSetError(S.NO_FILES_SELECTED)
    missing = [p for p in paths if not os.path.isfile(p)]
    if missing:
        raise DocSetError(S.FILES_NOT_FOUND.format(names=", ".join(os.path.basename(p) for p in missing)))

    task = _classify(user_request, paths)

    if task in (TASK_MERGE, TASK_EXTRACT_MERGE, TASK_COMPARE, TASK_RECONCILE) and len(paths) < 2:
        return DocSetResult(
            task=task,
            analysis="Cần chọn ít nhất 2 file để so sánh/đối chiếu/gộp — hiện chỉ có 1 file "
                     "trong ngữ cảnh.")

    if task == TASK_MERGE:
        return _handle_merge(paths, user_request)

    files, skipped = _extract_all(paths)
    if not files:
        raise DocSetError(S.NO_READABLE_CONTENT)

    if task == TASK_EXTRACT_MERGE:
        result = _handle_extract_merge(user_request, files)
    elif task == TASK_SEARCH:
        result = _handle_search(files, user_request)
    elif task == TASK_CHECKLIST:
        result = _handle_checklist(files, user_request, checklist)
    elif task in (TASK_COMPARE, TASK_RECONCILE):
        result = _handle_compare(files, user_request, task)
    else:
        result = _handle_general(files, user_request, history, business)

    if skipped:
        result.analysis = (f"⚠ Bỏ qua {len(skipped)} file không đọc được: "
                           + ", ".join(skipped) + "\n\n" + result.analysis)
    return result
