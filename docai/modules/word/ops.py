"""The set of Word (.docx) edit operations for the edit-via-chat flow.

Two parts:
  · `outline()` — describes the document's structure (paragraphs, pages,
    tables, styles) sent to the server so Claude knows what to target.
  · `apply_edits()` — executes the list of edit commands Claude returns.

PAGE boundaries are determined by manual page breaks in the file (page break
or the page-break-before property) — not pages as re-laid-out by Word.
"""
import copy
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ...logging_config import get_logger, log_call

logger = get_logger(__name__)


class WordOpError(Exception):
    """Invalid edit command — message shown to the user in Vietnamese."""


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_LIST_STYLES = {"bullet": "List Bullet", "number": "List Number"}

# Allowed operations — also the list the server describes to Claude.
OPS = (
    "append_paragraph",     # add a paragraph at the end of the document
    "insert_paragraph",     # insert a paragraph before/after a specific one
    "append_to_paragraph",  # append text to the end of an existing paragraph
    "replace_paragraph",    # replace a paragraph's entire content
    "delete_paragraph",     # delete a paragraph
    "replace_text",         # find & replace a string across the whole document
    "format_paragraph",     # font/size/bold/italic/underline/color/align/indent/line-spacing…
    "set_heading",          # turn a paragraph into a level-N heading
    "set_list",             # set/unset a bullet or numbered list
    "add_page",             # add a new page at the end
    "insert_page",          # insert a new page after a specific one
    "delete_page",          # delete a page
    "create_table",         # create a new table
    "add_table_row",        # add a row to an existing table
    "delete_table_row",     # delete a row from a table
    "add_table_column",     # add a column to an existing table
    "delete_table_column",  # delete a column from a table
    "set_table_cell",       # reset the content of an existing table cell
    "merge_table_cells",    # merge a range of table cells
    "split_table_cell",     # split a merged cell into multiple cells
    "format_table",         # borders & shading for a table/row/column/cell
)


# ── Read structure ────────────────────────────────────────────────────────

@log_call
def _has_page_break(para: Paragraph) -> bool:
    for br in para._p.findall(f".//{qn('w:br')}"):
        if br.get(qn("w:type")) == "page":
            return True
    return False


@log_call
def _page_numbers(doc) -> list[int]:
    """Page number for each paragraph (1-based), based on manual page breaks."""
    pages, page = [], 1
    for para in doc.paragraphs:
        if para.paragraph_format.page_break_before:
            page += 1
        pages.append(page)
        if _has_page_break(para):
            page += 1          # the next paragraph lands on the following page
    return pages


@log_call
def outline(path: str, max_chars: int = 90) -> dict:
    """The document's structure to send to the server as context for Claude."""
    doc = DocxDocument(path)
    pages = _page_numbers(doc)
    items = []
    for para_index, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        preview = text[:max_chars] + ("…" if len(text) > max_chars else "")
        items.append({
            "index": para_index,
            "page": pages[para_index],
            "style": para.style.name if para.style else "Normal",
            "text": preview,
            "page_break_after": _has_page_break(para),
        })

    tables = []
    for ti, table in enumerate(doc.tables):
        header = [cell.text.strip()[:30] for cell in table.rows[0].cells] if table.rows else []
        tables.append({
            "index": ti, "rows": len(table.rows), "cols": len(table.columns),
            "header": header,
        })

    return {
        "paragraphs": items,
        "page_count": max(pages) if pages else 1,
        "table_count": len(doc.tables),
        "tables": tables,
    }


@log_call
def outline_text(path: str) -> str:
    """A compact text-form outline for the prompt."""
    data = outline(path)
    lines = [f"Tài liệu có {len(data['paragraphs'])} đoạn, "
             f"{data['page_count']} trang, {data['table_count']} bảng."]
    for table_info in data["tables"]:
        header = " | ".join(header_cell for header_cell in table_info["header"] if header_cell) or "(trống)"
        lines.append(f"[Bảng {table_info['index']}] {table_info['rows']}×{table_info['cols']}, hàng đầu: {header}")
    current_page = 0
    for it in data["paragraphs"]:
        if it["page"] != current_page:
            current_page = it["page"]
            lines.append(f"--- Trang {current_page} ---")
        style = "" if it["style"] == "Normal" else f" [{it['style']}]"
        body = it["text"] or "(đoạn trống)"
        lines.append(f"[{it['index']}]{style} {body}")
        if it["page_break_after"]:
            lines.append("(ngắt trang)")
    return "\n".join(lines)


# ── XML manipulation helpers ─────────────────────────────────────────────────

@log_call
def _insert_after(para: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    if style:
        _set_style(new_para, style)
    if text:
        new_para.add_run(text)
    return new_para


@log_call
def _set_style(para: Paragraph, style: str):
    try:
        para.style = style
    except KeyError:
        raise WordOpError(f"Tài liệu không có style «{style}».")


@log_call
def _set_text(para: Paragraph, text: str):
    """Replace the paragraph's content, keeping the first run's formatting."""
    runs = para.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        para.add_run(text)


@log_call
def _delete(para: Paragraph):
    para._p.getparent().remove(para._p)


@log_call
def _add_lines(after: Paragraph | None, doc, text: str,
               style: str | None = None) -> Paragraph:
    """Add text (possibly multi-line) — continuing after `after`, or at the
    end of the document."""
    last = after
    for line in text.split("\n"):
        if last is None:
            last = doc.add_paragraph(line)
            if style:
                _set_style(last, style)
        else:
            last = _insert_after(last, line, style)
    return last


# ── Execute edit commands ────────────────────────────────────────────────────

@log_call
def _need(edit: dict, key: str):
    val = edit.get(key)
    if val is None or (isinstance(val, str) and not val.strip()):
        raise WordOpError(f"Lệnh «{edit.get('op')}» thiếu thông tin: {key}.")
    return val


@log_call
def _para_at(paras: list[Paragraph], edit: dict) -> Paragraph:
    idx = _need(edit, "index")
    try:
        idx = int(idx)
    except (TypeError, ValueError):
        raise WordOpError(f"Vị trí đoạn không hợp lệ: {idx}.")
    if not 0 <= idx < len(paras):
        raise WordOpError(
            f"Tài liệu chỉ có {len(paras)} đoạn (0–{len(paras) - 1}), "
            f"không có đoạn {idx}.")
    return paras[idx]


@log_call
def _page_paras(paras: list[Paragraph], pages: list[int],
                page: int) -> list[Paragraph]:
    try:
        page = int(page)
    except (TypeError, ValueError):
        raise WordOpError(f"Số trang không hợp lệ: {page}.")
    found = [para for para, pg in zip(paras, pages) if pg == page]
    if not found:
        total = max(pages) if pages else 1
        raise WordOpError(f"Tài liệu chỉ có {total} trang, không có trang {page}.")
    return found


@log_call
def _table_at(doc, edit: dict):
    idx = _need(edit, "table_index")
    try:
        idx = int(idx)
    except (TypeError, ValueError):
        raise WordOpError(f"Chỉ số bảng không hợp lệ: {idx}.")
    tables = doc.tables
    if not 0 <= idx < len(tables):
        raise WordOpError(
            f"Tài liệu chỉ có {len(tables)} bảng (0–{len(tables) - 1}), "
            f"không có bảng {idx}.")
    return tables[idx]


@log_call
def apply_edits(path: str, edits: list[dict]) -> list[str]:
    """Apply the edit commands to the file in order, save once. Returns a
    description of each command.

    Every position (index/page) is computed against the ORIGINAL document:
    paragraphs are held by XML reference before editing, so insert/delete
    doesn't shift the position of later commands.
    """
    if not edits:
        return []

    doc = DocxDocument(path)
    paras = list(doc.paragraphs)
    pages = _page_numbers(doc)
    notes: list[str] = []

    for edit in edits:
        op = (edit.get("op") or "").strip()
        if op not in OPS:
            raise WordOpError(f"Thao tác chưa hỗ trợ: «{op}».")
        notes.append(_run_op(doc, paras, pages, op, edit))

    try:
        doc.save(path)
    except PermissionError:
        raise WordOpError(
            f"Không ghi được «{Path(path).name}» — file đang mở trong Word, "
            "hãy đóng lại rồi thử lại.")
    return notes


@log_call
def _run_op(doc, paras: list[Paragraph], pages: list[int],
            op: str, edit: dict) -> str:
    style = edit.get("style") or None

    if op == "append_paragraph":
        text = _need(edit, "text")
        _add_lines(paras[-1] if paras else None, doc, text, style)
        return f"thêm «{_short(text)}» vào cuối tài liệu"

    if op == "insert_paragraph":
        text = _need(edit, "text")
        anchor = _para_at(paras, edit)
        position = (edit.get("position") or "after").lower()
        if position == "before":
            new = anchor.insert_paragraph_before("")
            _add_lines(new, doc, text, style)
            _delete(new)
            where = "trước"
        else:
            _add_lines(anchor, doc, text, style)
            where = "sau"
        return f"chèn «{_short(text)}» {where} đoạn {edit['index']}"

    if op == "append_to_paragraph":
        text = _need(edit, "text")
        para = _para_at(paras, edit)
        para.add_run(text)
        return f"nối «{_short(text)}» vào đoạn {edit['index']}"

    if op == "replace_paragraph":
        text = _need(edit, "text")
        para = _para_at(paras, edit)
        _set_text(para, text)
        return f"thay nội dung đoạn {edit['index']} thành «{_short(text)}»"

    if op == "delete_paragraph":
        para = _para_at(paras, edit)
        preview = _short(para.text) or "đoạn trống"
        _delete(para)
        return f"xóa đoạn {edit['index']} («{preview}»)"

    if op == "replace_text":
        find = _need(edit, "find")
        repl = edit.get("replace") or ""
        count = _replace_text(doc, find, repl)
        if not count:
            raise WordOpError(f"Không tìm thấy «{find}» trong tài liệu.")
        action = "xóa" if not repl else f"thay bằng «{_short(repl)}»"
        return f"tìm «{_short(find)}» ({count} chỗ) và {action}"

    if op == "format_paragraph":
        para = _para_at(paras, edit)
        return _format(para, edit)

    if op == "set_heading":
        para = _para_at(paras, edit)
        level = int(edit.get("level") or 1)
        if not 1 <= level <= 9:
            raise WordOpError("Cấp tiêu đề phải từ 1 đến 9.")
        _set_style(para, f"Heading {level}")
        return f"đặt đoạn {edit['index']} thành tiêu đề cấp {level}"

    if op == "set_list":
        para = _para_at(paras, edit)
        return _set_list(para, edit)

    if op == "add_page":
        text = _need(edit, "text")
        last = paras[-1] if paras else doc.add_paragraph("")
        brk = _insert_after(last)
        brk.add_run().add_break(WD_BREAK.PAGE)
        _add_lines(brk, doc, text, style)
        return f"thêm trang mới ở cuối với «{_short(text)}»"

    if op == "insert_page":
        text = _need(edit, "text")
        after_page = int(edit.get("after_page") or (max(pages) if pages else 1))
        last = _page_paras(paras, pages, after_page)[-1]
        if _has_page_break(last):
            # This page already ends with a page break → reuse that break,
            # then place a new break AFTER the content to push the old page
            # down. Inserting a break before the content would produce a
            # blank page.
            end = _add_lines(last, doc, text, style)
            brk = _insert_after(end)
            brk.add_run().add_break(WD_BREAK.PAGE)
        else:
            brk = _insert_after(last)
            brk.add_run().add_break(WD_BREAK.PAGE)
            _add_lines(brk, doc, text, style)
        return f"chèn trang mới sau trang {after_page} với «{_short(text)}»"

    if op == "create_table":
        return _create_table(doc, paras, edit)

    if op == "add_table_row":
        table = _table_at(doc, edit)
        return _add_table_row(table, edit)

    if op == "delete_table_row":
        table = _table_at(doc, edit)
        return _delete_table_row(table, edit)

    if op == "add_table_column":
        table = _table_at(doc, edit)
        return _add_table_column(table, edit)

    if op == "delete_table_column":
        table = _table_at(doc, edit)
        return _delete_table_column(table, edit)

    if op == "set_table_cell":
        table = _table_at(doc, edit)
        return _set_table_cell(table, edit)

    if op == "merge_table_cells":
        table = _table_at(doc, edit)
        return _merge_table_cells(table, edit)

    if op == "split_table_cell":
        table = _table_at(doc, edit)
        return _split_table_cell(table, edit)

    if op == "format_table":
        table = _table_at(doc, edit)
        return _format_table(table, edit)

    # delete_page — the last remaining op in OPS.
    page = int(_need(edit, "page"))
    victims = _page_paras(paras, pages, page)
    total = max(pages) if pages else 1
    for para in victims:
        _delete(para)
    # Deleting the last page → also clean up the now-orphaned page break at
    # the end of the previous page, to avoid a blank page.
    if page == total and page > 1:
        prev = _page_paras(paras, pages, page - 1)[-1]
        if _has_page_break(prev):
            if prev.text.strip():
                for br in prev._p.findall(f".//{qn('w:br')}"):
                    if br.get(qn("w:type")) == "page":
                        br.getparent().remove(br)
            else:
                _delete(prev)
    return f"xóa trang {page} ({len(victims)} đoạn)"


@log_call
def _is_simple_run(run: Run) -> bool:
    """True if the run contains exactly one <w:t> (no tab/break/multiple text
    runs) — safe to split in two. A run created by add_run() is always in
    this form."""
    run_elem = run._r
    children = [child for child in run_elem if child.tag != qn("w:rPr")]
    return len(children) == 1 and children[0].tag == qn("w:t")


@log_call
def _split_run(run: Run, offset: int) -> tuple[Run, Run]:
    """Split a run into 2 runs at `offset` within run.text, the new run keeps
    the original run's formatting (rPr). Only use when `_is_simple_run(run)`."""
    text = run.text
    run_elem = run._r
    rPr = run_elem.find(qn("w:rPr"))
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(copy.deepcopy(rPr))
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = text[offset:]
    new_r.append(new_t)
    run_elem.addnext(new_r)
    run_elem.find(qn("w:t")).text = text[:offset]
    return Run(run_elem, run._parent), Run(new_r, run._parent)


@log_call
def _format_run_range(para: Paragraph, start: int, end: int) -> list[Run]:
    """Split any runs crossed by [start, end) then return exactly the runs
    that land fully within that range — so formatting only applies to the
    selected text, not the whole paragraph. A "complex" run (containing a
    tab/line-break) that gets crossed is skipped instead of split, to avoid
    corrupting its content."""
    runs = list(para.runs)
    changed = True
    while changed:
        changed = False
        pos = 0
        for run_index, run in enumerate(runs):
            rlen = len(run.text)
            rstart, rend = pos, pos + rlen
            cut = None
            if rstart < start < rend:
                cut = start - rstart
            elif rstart < end < rend:
                cut = end - rstart
            if cut is not None:
                if _is_simple_run(run):
                    runs[run_index:run_index + 1] = list(_split_run(run, cut))
                    changed = True
                    break
                pos = rend
                continue
            pos = rend

    result = []
    pos = 0
    for run in runs:
        rlen = len(run.text)
        if rlen and pos >= start and pos + rlen <= end:
            result.append(run)
        pos += rlen
    return result


@log_call
def _set_run_font(run: Run, name: str):
    """Change the font, setting both ascii/hAnsi (via the API) and
    eastAsia/cs (via raw XML) so Word doesn't silently swap the font for
    some character ranges."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


@log_call
def _clean_hex(color) -> str:
    hex_color = str(color).strip().lstrip("#").upper()
    if len(hex_color) != 6 or any(ch not in "0123456789ABCDEF" for ch in hex_color):
        raise WordOpError(f"Mã màu không hợp lệ: {color}.")
    return hex_color


@log_call
def _format(para: Paragraph, edit: dict) -> str:
    changes = []
    target = edit.get("text") or None
    scope = f"đoạn {edit['index']}"

    if target:
        start = para.text.find(target)
        if start < 0:
            raise WordOpError(
                f"Không tìm thấy «{_short(target)}» trong đoạn {edit['index']}.")
        runs = _format_run_range(para, start, start + len(target))
        if not runs:
            raise WordOpError(
                f"Không định dạng được «{_short(target)}» trong đoạn "
                f"{edit['index']} — vùng chữ này quá phức tạp (chứa tab/ngắt dòng).")
        scope = f"«{_short(target)}» trong đoạn {edit['index']}"
    else:
        runs = para.runs or [para.add_run("")]

    for key, label in (("bold", "in đậm"), ("italic", "in nghiêng"),
                       ("underline", "gạch chân")):
        if edit.get(key) is not None:
            val = bool(edit[key])
            for run in runs:
                setattr(run, key, val)
            changes.append(label if val else f"bỏ {label}")

    if edit.get("size"):
        size = float(edit["size"])
        for run in runs:
            run.font.size = Pt(size)
        changes.append(f"cỡ chữ {size:g}pt")

    if edit.get("font_name"):
        name = str(edit["font_name"])
        for run in runs:
            _set_run_font(run, name)
        changes.append(f"font «{name}»")

    if edit.get("font_color"):
        hex_color = _clean_hex(edit["font_color"])
        for run in runs:
            run.font.color.rgb = RGBColor.from_string(hex_color)
        changes.append(f"màu chữ #{hex_color}")

    if edit.get("align"):
        align = str(edit["align"]).lower()
        if align not in _ALIGN:
            raise WordOpError(f"Kiểu căn lề chưa hỗ trợ: {align}.")
        para.alignment = _ALIGN[align]
        changes.append({"left": "căn trái", "center": "căn giữa",
                        "right": "căn phải", "justify": "căn đều"}[align])

    if edit.get("indent_left") is not None:
        para.paragraph_format.left_indent = Cm(float(edit["indent_left"]))
        changes.append(f"thụt lề trái {edit['indent_left']:g}cm")

    if edit.get("indent_right") is not None:
        para.paragraph_format.right_indent = Cm(float(edit["indent_right"]))
        changes.append(f"thụt lề phải {edit['indent_right']:g}cm")

    if edit.get("indent_first_line") is not None:
        para.paragraph_format.first_line_indent = Cm(float(edit["indent_first_line"]))
        val = float(edit["indent_first_line"])
        changes.append(f"thụt dòng đầu {val:g}cm" if val >= 0 else f"thụt treo {-val:g}cm")

    if edit.get("line_spacing") is not None:
        para.paragraph_format.line_spacing = float(edit["line_spacing"])
        changes.append(f"giãn dòng {float(edit['line_spacing']):g}")

    if edit.get("space_before") is not None:
        para.paragraph_format.space_before = Pt(float(edit["space_before"]))
        changes.append(f"khoảng cách trước đoạn {edit['space_before']:g}pt")

    if edit.get("space_after") is not None:
        para.paragraph_format.space_after = Pt(float(edit["space_after"]))
        changes.append(f"khoảng cách sau đoạn {edit['space_after']:g}pt")

    if edit.get("style"):
        _set_style(para, str(edit["style"]))
        changes.append(f"style «{edit['style']}»")

    if not changes:
        raise WordOpError("Lệnh «format_paragraph» không nêu định dạng nào.")
    return f"định dạng {scope}: {', '.join(changes)}"


@log_call
def _set_list(para: Paragraph, edit: dict) -> str:
    list_type = str(edit.get("list_type") or "").strip().lower()
    if list_type == "none":
        _set_style(para, "Normal")
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                pPr.remove(numPr)
        return f"bỏ danh sách ở đoạn {edit['index']}"

    style_name = _LIST_STYLES.get(list_type)
    if not style_name:
        raise WordOpError('list_type phải là "bullet", "number" hoặc "none".')
    _set_style(para, style_name)
    label = "gạch đầu dòng" if list_type == "bullet" else "đánh số"
    return f"đặt đoạn {edit['index']} thành danh sách {label}"


@log_call
def _replace_text(doc, find: str, repl: str) -> int:
    """Replace a string in every paragraph (including inside tables). Returns
    the number of paragraphs changed."""
    count = 0

    def _in_paragraphs(paragraphs):
        nonlocal count
        for para in paragraphs:
            if find in para.text:
                _set_text(para, para.text.replace(find, repl))
                count += 1

    _in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _in_paragraphs(cell.paragraphs)
    return count


# ── Tables ───────────────────────────────────────────────────────────────────

@log_call
def _create_table(doc, paras: list[Paragraph], edit: dict) -> str:
    rows = int(_need(edit, "rows"))
    cols = int(_need(edit, "cols"))
    if rows < 1 or cols < 1:
        raise WordOpError("Số hàng/cột của bảng phải lớn hơn 0.")
    data = edit.get("data")  # optional list[list[str]] — pre-fills content

    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        logger.debug("'Table Grid' style not in document — table created without default borders.")
        # style isn't available in the document — table still gets created, just without default borders

    if data:
        for row_idx, row_vals in enumerate(data):
            if row_idx >= rows:
                break
            for col_idx, val in enumerate(row_vals):
                if col_idx >= cols:
                    break
                table.cell(row_idx, col_idx).text = str(val)

    after_index = edit.get("after_index")
    if after_index is not None:
        try:
            ai_idx = int(after_index)
        except (TypeError, ValueError):
            raise WordOpError(f"Vị trí chèn bảng không hợp lệ: {after_index}.")
        if not 0 <= ai_idx < len(paras):
            raise WordOpError(
                f"Tài liệu chỉ có {len(paras)} đoạn (0–{len(paras) - 1}), "
                f"không có đoạn {ai_idx}.")
        paras[ai_idx]._p.addnext(table._tbl)
        where = f"sau đoạn {ai_idx}"
    else:
        where = "cuối tài liệu"

    return f"tạo bảng {rows}×{cols} ở {where}"


@log_call
def _add_table_row(table, edit: dict) -> str:
    data = edit.get("data")
    original_count = len(table.rows)
    row = table.add_row()
    if data:
        for col_idx, val in enumerate(data):
            if col_idx >= len(row.cells):
                break
            row.cells[col_idx].text = str(val)

    at = edit.get("at")
    if at is not None:
        try:
            at = int(at)
        except (TypeError, ValueError):
            raise WordOpError(f"Vị trí chèn hàng không hợp lệ: {at}.")
        if not 0 <= at <= original_count:
            raise WordOpError(
                f"Bảng có {original_count} hàng — không chèn được ở vị trí {at}.")
        if at < original_count:
            table.rows[at]._tr.addprevious(row._tr)
            return f"thêm hàng mới vào bảng ở vị trí {at}"
    return "thêm hàng mới vào cuối bảng"


@log_call
def _delete_table_row(table, edit: dict) -> str:
    row_index = _need(edit, "row_index")
    try:
        row_index = int(row_index)
    except (TypeError, ValueError):
        raise WordOpError(f"Chỉ số hàng không hợp lệ: {row_index}.")
    rows = table.rows
    if not 0 <= row_index < len(rows):
        raise WordOpError(f"Bảng chỉ có {len(rows)} hàng, không có hàng {row_index}.")
    if len(rows) == 1:
        raise WordOpError("Không thể xóa hàng cuối cùng của bảng.")
    rows[row_index]._tr.getparent().remove(rows[row_index]._tr)
    return f"xóa hàng {row_index} của bảng"


@log_call
def _table_grid(table):
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblPr.addnext(grid)
        else:
            tbl.insert(0, grid)
    return grid


@log_call
def _add_table_column(table, edit: dict) -> str:
    data = edit.get("data")
    at = edit.get("at")
    col_count = len(table.columns)
    if at is None:
        at = col_count
    else:
        try:
            at = int(at)
        except (TypeError, ValueError):
            raise WordOpError(f"Vị trí chèn cột không hợp lệ: {at}.")
        if not 0 <= at <= col_count:
            raise WordOpError(f"Bảng có {col_count} cột — không chèn được ở vị trí {at}.")

    grid = _table_grid(table)
    grid_cols = grid.findall(qn("w:gridCol"))
    ref_width = grid_cols[0].get(qn("w:w")) if grid_cols else "2000"
    new_grid_col = OxmlElement("w:gridCol")
    new_grid_col.set(qn("w:w"), ref_width)
    if grid_cols and at < len(grid_cols):
        grid_cols[at].addprevious(new_grid_col)
    else:
        grid.append(new_grid_col)

    for row_idx, row in enumerate(table.rows):
        tr = row._tr
        cells = tr.findall(qn("w:tc"))
        template = cells[0] if cells else None
        new_tc = copy.deepcopy(template) if template is not None else OxmlElement("w:tc")
        # Clear the old content, keep the cell formatting (tcPr) copied from the template cell in the same row.
        for para_elem in new_tc.findall(qn("w:p")):
            new_tc.remove(para_elem)
        new_p = OxmlElement("w:p")
        new_tc.append(new_p)
        if data and row_idx < len(data) and str(data[row_idx]):
            Paragraph(new_p, row.cells[0]._parent).add_run(str(data[row_idx]))

        if cells and at < len(cells):
            cells[at].addprevious(new_tc)
        else:
            tr.append(new_tc)

    return f"thêm cột mới vào bảng ở vị trí {at}"


@log_call
def _delete_table_column(table, edit: dict) -> str:
    col_index = _need(edit, "col_index")
    try:
        col_index = int(col_index)
    except (TypeError, ValueError):
        raise WordOpError(f"Chỉ số cột không hợp lệ: {col_index}.")
    col_count = len(table.columns)
    if not 0 <= col_index < col_count:
        raise WordOpError(f"Bảng chỉ có {col_count} cột, không có cột {col_index}.")
    if col_count == 1:
        raise WordOpError("Không thể xóa cột cuối cùng của bảng.")

    grid = _table_grid(table)
    grid_cols = grid.findall(qn("w:gridCol"))
    if col_index < len(grid_cols):
        grid.remove(grid_cols[col_index])

    # Simplification: assumes each row has exactly 1 cell per grid column (no
    # horizontally merged cells). A table with complex merges may need manual
    # cleanup after deletion.
    for row in table.rows:
        cells = row._tr.findall(qn("w:tc"))
        if col_index < len(cells):
            row._tr.remove(cells[col_index])

    return f"xóa cột {col_index} của bảng"


@log_call
def _set_table_cell(table, edit: dict) -> str:
    """Reset the entire text content of an EXISTING table cell — doesn't add/
    delete a row or column. Keeps the first run's formatting in the cell
    (same as how replace_paragraph keeps formatting when replacing a
    paragraph's content)."""
    row_idx = int(_need(edit, "row"))
    col_idx = int(_need(edit, "col"))
    text = _need(edit, "text")
    n_rows, n_cols = len(table.rows), len(table.columns)
    if not 0 <= row_idx < n_rows or not 0 <= col_idx < n_cols:
        raise WordOpError(f"Ô ({row_idx},{col_idx}) vượt quá kích thước bảng ({n_rows}×{n_cols}).")

    cell = table.cell(row_idx, col_idx)
    paras = cell.paragraphs
    first = paras[0] if paras else cell.add_paragraph()
    _set_text(first, str(text))
    for extra in paras[1:]:
        _delete(extra)
    return f"đặt ô ({row_idx},{col_idx}) của bảng thành «{_short(str(text))}»"


@log_call
def _merge_table_cells(table, edit: dict) -> str:
    r1 = int(_need(edit, "start_row"))
    c1 = int(_need(edit, "start_col"))
    r2 = int(_need(edit, "end_row"))
    c2 = int(_need(edit, "end_col"))
    n_rows, n_cols = len(table.rows), len(table.columns)
    for name, val, limit in (("start_row", r1, n_rows), ("end_row", r2, n_rows),
                             ("start_col", c1, n_cols), ("end_col", c2, n_cols)):
        if not 0 <= val < limit:
            raise WordOpError(f"{name} = {val} vượt quá kích thước bảng ({n_rows}×{n_cols}).")
    try:
        table.cell(r1, c1).merge(table.cell(r2, c2))
    except Exception as exc:  # noqa: BLE001 — unusual table structure error
        raise WordOpError(f"Không gộp được ô: {exc}")
    return f"gộp ô ({r1},{c1})–({r2},{c2}) trong bảng"


@log_call
def _ensure_tcPr(tc):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    return tcPr


@log_call
def _set_grid_span(tc, span: int):
    tcPr = _ensure_tcPr(tc)
    gs = tcPr.find(qn("w:gridSpan"))
    if span <= 1:
        if gs is not None:
            tcPr.remove(gs)
        return
    if gs is None:
        gs = OxmlElement("w:gridSpan")
        tcPr.insert(0, gs)
    gs.set(qn("w:val"), str(span))


@log_call
def _split_table_cell(table, edit: dict) -> str:
    """Split a merged cell (gridSpan > 1) into several smaller cells, dividing
    the occupied grid columns evenly. A never-merged cell (gridSpan == 1)
    can't be split — that requires adding a new column to the whole table
    first."""
    row_idx = int(_need(edit, "row"))
    col_idx = int(_need(edit, "col"))
    parts = int(edit.get("cols") or 2)
    if parts < 2:
        raise WordOpError("Số phần tách phải từ 2 trở lên.")
    n_rows, n_cols = len(table.rows), len(table.columns)
    if not 0 <= row_idx < n_rows or not 0 <= col_idx < n_cols:
        raise WordOpError(f"Ô ({row_idx},{col_idx}) vượt quá kích thước bảng ({n_rows}×{n_cols}).")

    tc = table.cell(row_idx, col_idx)._tc
    tcPr = tc.find(qn("w:tcPr"))
    gs_el = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
    span = int(gs_el.get(qn("w:val"))) if gs_el is not None else 1
    if span < parts:
        raise WordOpError(
            f"Ô này đang chiếm {span} cột lưới — không đủ để tách thành {parts} phần "
            "(chỉ tách được ô đã gộp từ nhiều ô trước đó).")

    for para_elem in tc.findall(qn("w:p"))[1:]:
        tc.remove(para_elem)
    first_p = tc.find(qn("w:p"))
    if first_p is not None:
        for child in list(first_p):
            if child.tag != qn("w:pPr"):
                first_p.remove(child)

    base, extra = divmod(span, parts)
    spans = [base + (1 if idx < extra else 0) for idx in range(parts)]

    _set_grid_span(tc, spans[0])
    anchor = tc
    for part_span in spans[1:]:
        new_tc = copy.deepcopy(tc)
        _set_grid_span(new_tc, part_span)
        anchor.addnext(new_tc)
        anchor = new_tc

    return f"tách ô ({row_idx},{col_idx}) thành {n} phần"


@log_call
def _cell_set_shading(cell, color) -> None:
    hex_color = _clean_hex(color)
    tcPr = _ensure_tcPr(cell._tc)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


@log_call
def _cell_set_borders(cell, color, size: int) -> None:
    hex_color = _clean_hex(color)
    tcPr = _ensure_tcPr(cell._tc)
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), hex_color)


@log_call
def _table_set_borders(table, color, size: int) -> None:
    hex_color = _clean_hex(color)
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), hex_color)


_VALIGN = {"top", "center", "bottom"}
_VALIGN_LABEL = {"top": "căn trên", "center": "căn giữa dọc", "bottom": "căn dưới"}
_FONT_KEYS = ("bold", "italic", "underline", "font_size", "font_name", "font_color")


@log_call
def _cell_set_valign(cell, valign: str) -> None:
    tcPr = _ensure_tcPr(cell._tc)
    va = tcPr.find(qn("w:vAlign"))
    if va is None:
        va = OxmlElement("w:vAlign")
        tcPr.append(va)
    va.set(qn("w:val"), valign)


@log_call
def _apply_cell_format(cell, edit: dict) -> None:
    """Apply shading/border/font/alignment to a cell — shared by every scope
    of format_table. `edit["border"]` is a flag already decided upstream by
    _format_table (scope=table uses its own table-level border, not repeated
    here)."""
    if edit.get("shading_color"):
        _cell_set_shading(cell, edit["shading_color"])
    if edit.get("border"):
        _cell_set_borders(cell, edit.get("border_color") or "000000",
                          int(edit.get("border_size") or 4))

    if any(edit.get(font_key) is not None for font_key in _FONT_KEYS):
        for para in cell.paragraphs:
            runs = para.runs or [para.add_run("")]
            for key, label in (("bold", "bold"), ("italic", "italic"), ("underline", "underline")):
                if edit.get(key) is not None:
                    val = bool(edit[key])
                    for run in runs:
                        setattr(run, label, val)
            if edit.get("font_size"):
                size = float(edit["font_size"])
                for run in runs:
                    run.font.size = Pt(size)
            if edit.get("font_name"):
                name = str(edit["font_name"])
                for run in runs:
                    _set_run_font(run, name)
            if edit.get("font_color"):
                hex_color = _clean_hex(edit["font_color"])
                for run in runs:
                    run.font.color.rgb = RGBColor.from_string(hex_color)

    if edit.get("align"):
        align = str(edit["align"]).lower()
        if align not in _ALIGN:
            raise WordOpError(f"Kiểu căn ngang chưa hỗ trợ: {align}.")
        for para in cell.paragraphs:
            para.alignment = _ALIGN[align]

    if edit.get("valign"):
        valign = str(edit["valign"]).lower()
        if valign not in _VALIGN:
            raise WordOpError(f"Kiểu căn dọc chưa hỗ trợ: {valign}.")
        _cell_set_valign(cell, valign)


@log_call
def _format_table(table, edit: dict) -> str:
    scope = str(edit.get("scope") or "table").lower()
    color = edit.get("shading_color")
    border = edit.get("border")
    has_font = any(edit.get(font_key) is not None for font_key in _FONT_KEYS)
    has_align = edit.get("align") is not None
    has_valign = edit.get("valign") is not None

    if not (color or border or has_font or has_align or has_valign):
        raise WordOpError(
            "Lệnh «format_table» cần ít nhất một thuộc tính: shading_color, "
            "border, hoặc font/căn lề (bold, italic, underline, font_size, "
            "font_name, font_color, align, valign).")

    n_rows, n_cols = len(table.rows), len(table.columns)
    cell_border = border

    if scope == "table":
        cells = [cell for row in table.rows for cell in row.cells]
        if border:
            _table_set_borders(table, edit.get("border_color") or "000000",
                               int(edit.get("border_size") or 4))
            cell_border = False  # table-level border already handled — don't repeat it per cell
        where = "cả bảng"
    elif scope == "row":
        row_idx = int(_need(edit, "row"))
        if not 0 <= row_idx < n_rows:
            raise WordOpError(f"Bảng chỉ có {n_rows} hàng, không có hàng {row_idx}.")
        cells = list(table.rows[row_idx].cells)
        where = f"hàng {row_idx} của bảng"
    elif scope == "col":
        col_idx = int(_need(edit, "col"))
        if not 0 <= col_idx < n_cols:
            raise WordOpError(f"Bảng chỉ có {n_cols} cột, không có cột {col_idx}.")
        cells = [row.cells[col_idx] for row in table.rows if col_idx < len(row.cells)]
        where = f"cột {col_idx} của bảng"
    elif scope == "cell":
        row_idx = int(_need(edit, "row"))
        col_idx = int(_need(edit, "col"))
        if not 0 <= row_idx < n_rows or not 0 <= col_idx < n_cols:
            raise WordOpError(f"Ô ({row_idx},{col_idx}) vượt quá kích thước bảng ({n_rows}×{n_cols}).")
        cells = [table.cell(row_idx, col_idx)]
        where = f"ô ({row_idx},{col_idx}) của bảng"
    else:
        raise WordOpError('scope phải là "table", "row", "col" hoặc "cell".')

    cell_edit = {**edit, "border": cell_border}
    for cell in cells:
        _apply_cell_format(cell, cell_edit)

    changes = []
    if color:
        changes.append(f"nền #{_clean_hex(color)}")
    if border:
        changes.append("viền")
    if has_font:
        changes.append("font/kiểu chữ")
    if has_align:
        changes.append(f"căn ngang {edit['align']}")
    if has_valign:
        changes.append(f"căn dọc {_VALIGN_LABEL[str(edit['valign']).lower()]}")
    return f"định dạng {where}: {', '.join(changes)}"


@log_call
def _short(text: str, limit: int = 40) -> str:
    text = " ".join(str(text).split())
    return text if len(text) <= limit else text[:limit] + "…"
