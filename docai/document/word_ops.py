"""Bộ thao tác chỉnh sửa file Word (.docx) cho luồng sửa-bằng-chat.

Hai phần:
  · `outline()` — mô tả cấu trúc tài liệu (đoạn, trang, style) gửi lên server
    để Claude biết nhắm vào đâu.
  · `apply_edits()` — thực thi danh sách lệnh sửa Claude trả về.

Ranh giới TRANG được xác định theo ngắt trang thủ công trong file (page break
hoặc thuộc tính page-break-before) — không phải trang do Word dàn trang lại.
"""
import copy
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run


class WordOpError(Exception):
    """Lệnh sửa không hợp lệ — thông báo tiếng Việt cho người dùng."""


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Các thao tác được phép — cũng là danh sách server mô tả cho Claude.
OPS = (
    "append_paragraph",     # thêm đoạn ở cuối tài liệu
    "insert_paragraph",     # chèn đoạn trước/sau một đoạn cụ thể
    "append_to_paragraph",  # nối thêm chữ vào cuối một đoạn có sẵn
    "replace_paragraph",    # thay toàn bộ nội dung một đoạn
    "delete_paragraph",     # xóa một đoạn
    "replace_text",         # tìm & thay chuỗi trong toàn tài liệu
    "format_paragraph",     # in đậm/nghiêng/gạch chân/cỡ chữ/căn lề
    "set_heading",          # đặt đoạn thành tiêu đề cấp N
    "add_page",             # thêm trang mới ở cuối
    "insert_page",          # chèn trang mới sau một trang cụ thể
    "delete_page",          # xóa một trang
)


# ── Đọc cấu trúc ───────────────────────────────────────────────────────────

def _has_page_break(para: Paragraph) -> bool:
    for br in para._p.findall(f".//{qn('w:br')}"):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _page_numbers(doc) -> list[int]:
    """Số trang của từng đoạn (1-based), theo ngắt trang thủ công."""
    pages, page = [], 1
    for para in doc.paragraphs:
        if para.paragraph_format.page_break_before:
            page += 1
        pages.append(page)
        if _has_page_break(para):
            page += 1          # đoạn sau nằm ở trang kế tiếp
    return pages


def outline(path: str, max_chars: int = 90) -> dict:
    """Cấu trúc tài liệu để gửi lên server làm ngữ cảnh cho Claude."""
    doc = DocxDocument(path)
    pages = _page_numbers(doc)
    items = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        preview = text[:max_chars] + ("…" if len(text) > max_chars else "")
        items.append({
            "index": i,
            "page": pages[i],
            "style": para.style.name if para.style else "Normal",
            "text": preview,
            "page_break_after": _has_page_break(para),
        })
    return {
        "paragraphs": items,
        "page_count": max(pages) if pages else 1,
        "table_count": len(doc.tables),
    }


def outline_text(path: str) -> str:
    """Outline dạng văn bản gọn cho prompt."""
    data = outline(path)
    lines = [f"Tài liệu có {len(data['paragraphs'])} đoạn, "
             f"{data['page_count']} trang, {data['table_count']} bảng."]
    current_page = 0
    for it in data["paragraphs"]:
        if it["page"] != current_page:
            current_page = it["page"]
            lines.append(f"--- Trang {current_page} ---")
        style = "" if it["style"] == "Normal" else f" [{it['style']}]"
        body = it["text"] or "(đoạn trống)"
        lines.append(f"[{it['index']}]{style} {body}")
        if it["page_break_after"]:
            lines.append("(ngắt trang)")
    return "\n".join(lines)


# ── Tiện ích thao tác trên XML ─────────────────────────────────────────────

def _insert_after(para: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    if style:
        _set_style(new_para, style)
    if text:
        new_para.add_run(text)
    return new_para


def _set_style(para: Paragraph, style: str):
    try:
        para.style = style
    except KeyError:
        raise WordOpError(f"Tài liệu không có style «{style}».")


def _set_text(para: Paragraph, text: str):
    """Thay nội dung đoạn, giữ định dạng của run đầu tiên."""
    runs = para.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        para.add_run(text)


def _delete(para: Paragraph):
    para._p.getparent().remove(para._p)


def _add_lines(after: Paragraph | None, doc, text: str,
               style: str | None = None) -> Paragraph:
    """Thêm text (có thể nhiều dòng) — nối tiếp sau `after`, hoặc cuối tài liệu."""
    last = after
    for line in text.split("\n"):
        if last is None:
            last = doc.add_paragraph(line)
            if style:
                _set_style(last, style)
        else:
            last = _insert_after(last, line, style)
    return last


# ── Thực thi lệnh sửa ──────────────────────────────────────────────────────

def _need(edit: dict, key: str):
    val = edit.get(key)
    if val is None or (isinstance(val, str) and not val.strip()):
        raise WordOpError(f"Lệnh «{edit.get('op')}» thiếu thông tin: {key}.")
    return val


def _para_at(paras: list[Paragraph], edit: dict) -> Paragraph:
    idx = _need(edit, "index")
    try:
        idx = int(idx)
    except (TypeError, ValueError):
        raise WordOpError(f"Vị trí đoạn không hợp lệ: {idx}.")
    if not 0 <= idx < len(paras):
        raise WordOpError(
            f"Tài liệu chỉ có {len(paras)} đoạn (0–{len(paras) - 1}), "
            f"không có đoạn {idx}.")
    return paras[idx]


def _page_paras(paras: list[Paragraph], pages: list[int],
                page: int) -> list[Paragraph]:
    try:
        page = int(page)
    except (TypeError, ValueError):
        raise WordOpError(f"Số trang không hợp lệ: {page}.")
    found = [p for p, pg in zip(paras, pages) if pg == page]
    if not found:
        total = max(pages) if pages else 1
        raise WordOpError(f"Tài liệu chỉ có {total} trang, không có trang {page}.")
    return found


def apply_edits(path: str, edits: list[dict]) -> list[str]:
    """Áp lần lượt các lệnh sửa lên file, lưu một lần. Trả về mô tả từng lệnh.

    Mọi vị trí (index/page) đều tính theo tài liệu GỐC: các đoạn được nắm giữ
    bằng tham chiếu XML trước khi sửa, nên chèn/xóa không làm lệch vị trí của
    những lệnh sau.
    """
    if not edits:
        return []

    doc = DocxDocument(path)
    paras = list(doc.paragraphs)
    pages = _page_numbers(doc)
    notes: list[str] = []

    for edit in edits:
        op = (edit.get("op") or "").strip()
        if op not in OPS:
            raise WordOpError(f"Thao tác chưa hỗ trợ: «{op}».")
        notes.append(_run_op(doc, paras, pages, op, edit))

    try:
        doc.save(path)
    except PermissionError:
        raise WordOpError(
            f"Không ghi được «{Path(path).name}» — file đang mở trong Word, "
            "hãy đóng lại rồi thử lại.")
    return notes


def _run_op(doc, paras: list[Paragraph], pages: list[int],
            op: str, edit: dict) -> str:
    style = edit.get("style") or None

    if op == "append_paragraph":
        text = _need(edit, "text")
        _add_lines(paras[-1] if paras else None, doc, text, style)
        return f"thêm «{_short(text)}» vào cuối tài liệu"

    if op == "insert_paragraph":
        text = _need(edit, "text")
        anchor = _para_at(paras, edit)
        position = (edit.get("position") or "after").lower()
        if position == "before":
            new = anchor.insert_paragraph_before("")
            _add_lines(new, doc, text, style)
            _delete(new)
            where = "trước"
        else:
            _add_lines(anchor, doc, text, style)
            where = "sau"
        return f"chèn «{_short(text)}» {where} đoạn {edit['index']}"

    if op == "append_to_paragraph":
        text = _need(edit, "text")
        para = _para_at(paras, edit)
        para.add_run(text)
        return f"nối «{_short(text)}» vào đoạn {edit['index']}"

    if op == "replace_paragraph":
        text = _need(edit, "text")
        para = _para_at(paras, edit)
        _set_text(para, text)
        return f"thay nội dung đoạn {edit['index']} thành «{_short(text)}»"

    if op == "delete_paragraph":
        para = _para_at(paras, edit)
        preview = _short(para.text) or "đoạn trống"
        _delete(para)
        return f"xóa đoạn {edit['index']} («{preview}»)"

    if op == "replace_text":
        find = _need(edit, "find")
        repl = edit.get("replace") or ""
        count = _replace_text(doc, find, repl)
        if not count:
            raise WordOpError(f"Không tìm thấy «{find}» trong tài liệu.")
        action = "xóa" if not repl else f"thay bằng «{_short(repl)}»"
        return f"tìm «{_short(find)}» ({count} chỗ) và {action}"

    if op == "format_paragraph":
        para = _para_at(paras, edit)
        return _format(para, edit)

    if op == "set_heading":
        para = _para_at(paras, edit)
        level = int(edit.get("level") or 1)
        if not 1 <= level <= 9:
            raise WordOpError("Cấp tiêu đề phải từ 1 đến 9.")
        _set_style(para, f"Heading {level}")
        return f"đặt đoạn {edit['index']} thành tiêu đề cấp {level}"

    if op == "add_page":
        text = _need(edit, "text")
        last = paras[-1] if paras else doc.add_paragraph("")
        brk = _insert_after(last)
        brk.add_run().add_break(WD_BREAK.PAGE)
        _add_lines(brk, doc, text, style)
        return f"thêm trang mới ở cuối với «{_short(text)}»"

    if op == "insert_page":
        text = _need(edit, "text")
        after_page = int(edit.get("after_page") or (max(pages) if pages else 1))
        last = _page_paras(paras, pages, after_page)[-1]
        if _has_page_break(last):
            # Trang này đã có ngắt trang ở cuối → dùng luôn ngắt đó, rồi đặt
            # một ngắt mới SAU nội dung để đẩy trang cũ xuống. Nếu chèn thêm
            # ngắt trước nội dung sẽ sinh ra một trang trắng.
            end = _add_lines(last, doc, text, style)
            brk = _insert_after(end)
            brk.add_run().add_break(WD_BREAK.PAGE)
        else:
            brk = _insert_after(last)
            brk.add_run().add_break(WD_BREAK.PAGE)
            _add_lines(brk, doc, text, style)
        return f"chèn trang mới sau trang {after_page} với «{_short(text)}»"

    # delete_page
    page = int(_need(edit, "page"))
    victims = _page_paras(paras, pages, page)
    total = max(pages) if pages else 1
    for para in victims:
        _delete(para)
    # Xóa trang cuối → dọn nốt ngắt trang thừa ở cuối trang trước, tránh trang trắng.
    if page == total and page > 1:
        prev = _page_paras(paras, pages, page - 1)[-1]
        if _has_page_break(prev):
            if prev.text.strip():
                for br in prev._p.findall(f".//{qn('w:br')}"):
                    if br.get(qn("w:type")) == "page":
                        br.getparent().remove(br)
            else:
                _delete(prev)
    return f"xóa trang {page} ({len(victims)} đoạn)"


def _is_simple_run(run: Run) -> bool:
    """True nếu run chỉ chứa đúng một <w:t> (không tab/break/nhiều đoạn chữ)
    — an toàn để tách đôi. Run do add_run() tạo ra luôn ở dạng này."""
    r = run._r
    children = [c for c in r if c.tag != qn("w:rPr")]
    return len(children) == 1 and children[0].tag == qn("w:t")


def _split_run(run: Run, offset: int) -> tuple[Run, Run]:
    """Tách run thành 2 run tại vị trí `offset` trong run.text, run mới giữ
    nguyên định dạng (rPr) của run gốc. Chỉ dùng khi `_is_simple_run(run)`."""
    text = run.text
    r = run._r
    rPr = r.find(qn("w:rPr"))
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(copy.deepcopy(rPr))
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = text[offset:]
    new_r.append(new_t)
    r.addnext(new_r)
    r.find(qn("w:t")).text = text[:offset]
    return Run(r, run._parent), Run(new_r, run._parent)


def _format_run_range(para: Paragraph, start: int, end: int) -> list[Run]:
    """Tách các run bị cắt ngang bởi [start, end) rồi trả về đúng những run
    nằm gọn trong khoảng đó — để định dạng chỉ áp lên phần chữ đã chọn, không
    lan ra cả đoạn. Run "phức tạp" (chứa tab/ngắt dòng) bị cắt ngang thì được
    bỏ qua thay vì tách để tránh hỏng nội dung."""
    runs = list(para.runs)
    changed = True
    while changed:
        changed = False
        pos = 0
        for i, run in enumerate(runs):
            rlen = len(run.text)
            rstart, rend = pos, pos + rlen
            cut = None
            if rstart < start < rend:
                cut = start - rstart
            elif rstart < end < rend:
                cut = end - rstart
            if cut is not None:
                if _is_simple_run(run):
                    runs[i:i + 1] = list(_split_run(run, cut))
                    changed = True
                    break
                pos = rend
                continue
            pos = rend

    result = []
    pos = 0
    for run in runs:
        rlen = len(run.text)
        if rlen and pos >= start and pos + rlen <= end:
            result.append(run)
        pos += rlen
    return result


def _format(para: Paragraph, edit: dict) -> str:
    changes = []
    target = edit.get("text") or None
    scope = f"đoạn {edit['index']}"

    if target:
        start = para.text.find(target)
        if start < 0:
            raise WordOpError(
                f"Không tìm thấy «{_short(target)}» trong đoạn {edit['index']}.")
        runs = _format_run_range(para, start, start + len(target))
        if not runs:
            raise WordOpError(
                f"Không định dạng được «{_short(target)}» trong đoạn "
                f"{edit['index']} — vùng chữ này quá phức tạp (chứa tab/ngắt dòng).")
        scope = f"«{_short(target)}» trong đoạn {edit['index']}"
    else:
        runs = para.runs or [para.add_run("")]

    for key, label in (("bold", "in đậm"), ("italic", "in nghiêng"),
                       ("underline", "gạch chân")):
        if edit.get(key) is not None:
            val = bool(edit[key])
            for run in runs:
                setattr(run, key, val)
            changes.append(label if val else f"bỏ {label}")

    if edit.get("size"):
        size = float(edit["size"])
        for run in runs:
            run.font.size = Pt(size)
        changes.append(f"cỡ chữ {size:g}pt")

    if edit.get("align"):
        align = str(edit["align"]).lower()
        if align not in _ALIGN:
            raise WordOpError(f"Kiểu căn lề chưa hỗ trợ: {align}.")
        para.alignment = _ALIGN[align]
        changes.append({"left": "căn trái", "center": "căn giữa",
                        "right": "căn phải", "justify": "căn đều"}[align])

    if edit.get("style"):
        _set_style(para, str(edit["style"]))
        changes.append(f"style «{edit['style']}»")

    if not changes:
        raise WordOpError("Lệnh «format_paragraph» không nêu định dạng nào.")
    return f"định dạng {scope}: {', '.join(changes)}"


def _replace_text(doc, find: str, repl: str) -> int:
    """Thay chuỗi trong mọi đoạn (kể cả trong bảng). Trả về số đoạn bị đổi."""
    count = 0

    def _in_paragraphs(paragraphs):
        nonlocal count
        for para in paragraphs:
            if find in para.text:
                _set_text(para, para.text.replace(find, repl))
                count += 1

    _in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _in_paragraphs(cell.paragraphs)
    return count


def _short(text: str, limit: int = 40) -> str:
    text = " ".join(str(text).split())
    return text if len(text) <= limit else text[:limit] + "…"
